# PyQt5 Drag-and-Drop GUI Demo for MS Word Documents (demo.py)

import docx, os, sys, win32com.client
from shutil import copyfile
from PyQt5 import QtWidgets, QtCore

# pip install pyinstaller
# pip install python-docx
# pip install PyQt5
# pip install pypiwin32

# pyinstaller --onedir --onefile --windowed demo.py

def function1(readtext, destinationfile):
    
    # Insert your code here for analyzing readtext.
    if utf == True:
        outfile = open(destinationfile, 'w', encoding="utf-8")
    else:
        outfile = open(destinationfile, 'w')
    outfile.write('\nFunction 1:\n')
    outfile.write('Function 1 results are outputted here.')
    outfile.close()
    return


def function2(readtext, destinationfile):
    
    # Insert your code here for analyzing readtext.
    if utf == True:
        outfile = open(destinationfile, 'w', encoding="utf-8")
    else:
        outfile = open(destinationfile, 'w')
    outfile.write('\nFunction 2:\n')
    outfile.write('Function 2 results are outputted here.')
    outfile.close()
    return


def function3(readtext, destinationfile):
    
    # Insert your code here for analyzing readtext.
    if utf == True:
        outfile = open(destinationfile, 'w', encoding="utf-8")
    else:
        outfile = open(destinationfile, 'w')
    outfile.write('\nFunction 3:\n')
    outfile.write('Function 3 results are outputted here.')
    outfile.close()
    return

def function4(readtext, destinationfile):
    
    # Insert your code here for analyzing readtext.
    if utf == True:
        outfile = open(destinationfile, 'w', encoding="utf-8")
    else:
        outfile = open(destinationfile, 'w')
    outfile.write('\nFunction 4:\n')
    outfile.write('Function 4 results are outputted here.')
    outfile.close()
    return
    
def filenumber(destinationfile):
    for x in range(1,10000):
        filename = destinationfile + '(' + str(x) +').txt'
        # This function helps in assigning a unique number to a results file,
        # if there are multiple results files with the same name.
        # If SAMPLE_Function1(1).txt already exists, and SAMPLE_Function(2).txt does not
        # exist yet, then this function returns '2'.
        if os.access(filename, os.R_OK)==False: #Does a results file with that number not exist yet?
            return str(x)
    return str(0)

def opentextfile(targetfile):
    openfile = open(targetfile, 'r') # Open the target file in text format.
    global utf
    utf = False
    try:
        readtext = openfile.read() # Extract the entire text from the target file in ascii.
    except:
        openfile.close() # If that doesn't work, then extract entire text in utf-8.
        openfile = open(targetfile, 'r', encoding='utf-8')
        readtext = openfile.read()
        utf = True
    openfile.close()    # Close the file.
    return readtext

class MainApplication(QtWidgets.QWidget):
    def __init__(self):
        super(MainApplication, self).__init__()
        self.setObjectName("MainApplication")
        self.resize(370, 307)
        self.groupBox = QtWidgets.QGroupBox(self)
        self.groupBox.setGeometry(QtCore.QRect(20, 20, 331, 81))
        self.groupBox.setObjectName("groupBox")
        self.Function1 = QtWidgets.QRadioButton(self.groupBox)
        self.Function1.setGeometry(QtCore.QRect(40, 20, 111, 20))
        self.Function1.setChecked(True)
        self.Function1.setObjectName("Function1")
        self.Function2 = QtWidgets.QRadioButton(self.groupBox)
        self.Function2.setGeometry(QtCore.QRect(180, 20, 131, 20))
        self.Function2.setObjectName("Function2")
        self.Function3 = QtWidgets.QRadioButton(self.groupBox)
        self.Function3.setGeometry(QtCore.QRect(40, 50, 131, 20))
        self.Function3.setObjectName("Function3")
        self.Function4 = QtWidgets.QRadioButton(self.groupBox)
        self.Function4.setGeometry(QtCore.QRect(180, 50, 151, 20))
        self.Function4.setObjectName("Function4")
        self.pushButton = QtWidgets.QPushButton(self)
        self.pushButton.setGeometry(QtCore.QRect(20, 107, 331, 41))
        self.pushButton.setObjectName("pushButton")
        self.pushButton.clicked.connect(self.load_file_but)
        self.textBrowser = QtWidgets.QTextBrowser(self)
        self.textBrowser.setGeometry(QtCore.QRect(20, 160, 331, 131))
        self.textBrowser.setObjectName("textBrowser")
        self.setWindowTitle(QtWidgets.QApplication.translate("MainApplication", "Main Application", None))
        self.groupBox.setTitle(QtWidgets.QApplication.translate("MainApplication", "Select One:", None))
        self.Function1.setText(QtWidgets.QApplication.translate("MainApplication", "Function 1", None))
        self.Function2.setText(QtWidgets.QApplication.translate("MainApplication", "Function 2", None))
        self.Function3.setText(QtWidgets.QApplication.translate("MainApplication", "Function 3", None))
        self.Function4.setText(QtWidgets.QApplication.translate("MainApplication", "Function 4", None))
        self.pushButton.setText(QtWidgets.QApplication.translate("MainApplication", "Load File", None))
        self.textBrowser.setHtml(QtWidgets.QApplication.translate("MainApplication","<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:7.5pt; font-weight:400; font-style:normal;\">\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:8pt;\">Please close MS Word for Word 97-2003 (.doc) files</span></p>\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:8pt;\"><br /></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:12pt;\">Drag and Drop File Here</span></p></body></html>", None))

        # Enable dragging and dropping onto the GUI
        self.setAcceptDrops(True)
        self.show()

    def load_file_but(self):
        """
        Open a File dialog when the button is pressed
        :return:
        """

        # Get the file location
        self.fname = QtWidgets.QFileDialog.getOpenFileName(self, 'Open file')
        # Load the file from the location
        self.load_file()

    def load_file(self):
        targetfilename = self.fname[0].lower()
        copiedfile1 = False
        copiedfile2 = False
        if targetfilename[-4:len(targetfilename)]=='.doc': # Loading Word 97-2003 document using win32com.client
            try:
                wordapp = win32com.client.gencache.EnsureDispatch("Word.Application")
            except:
                self.errorpywin32() # If unsuccessful in loading the Word 97-2003 document, an error is generated.
            try:
                if targetfilename.find(' ')!=-1:
                    destfilename = targetfilename.replace(' ','_') 
                    # All spaces are replaced with underscore for the copy of the original.
                    copyfile(targetfilename, destfilename)
                    targetfilename = destfilename
                    copiedfile1 = True
                    # A copy of the Word document is made to be very sure that the original file is safe.
                else:
                    destfilename = targetfilename.replace('.doc','_.doc')
                    # If there are no spaces to be replaced, an underscore is appended to the end of the filename before the extension.
                    copyfile(targetfilename, destfilename)
                    targetfilename = destfilename
                    copiedfile1 = True
                    # A copy of the Word document is made to be very sure that the original file is safe.
                wordapp.Documents.Open(targetfilename)
                # The copy of the Word document is opened.
            except:
                self.errorFileload()
                # If the Word document cannot be opened for any reason, an error box will pop up.
                return
            targetfilename = targetfilename.replace('.doc','.txt') # The output file will have the txt extension.
            wordapp.ActiveDocument.SaveAs(targetfilename, FileFormat=win32com.client.constants.wdFormatTextLineBreaks)
            wordapp.ActiveWindow.Close()
            # The Word document is saved as a text document, and then the Word document is closed.
        if targetfilename[-5:len(targetfilename)]=='.docx': # Loading Word Document using python-docx.
            try:
                if targetfilename.find(' ')!=-1:
                    destfilename = targetfilename.replace(' ','_')
                    # All spaces are replaced with underscore for the copy of the original.
                    copyfile(targetfilename, destfilename)
                    targetfilename = destfilename
                    copiedfile2 = True
                    # A copy of the Word document is made to be very sure that the original file is safe.
                else:
                    destfilename = targetfilename.replace('.docx','_.docx')
                    # If there are no spaces to be replaced, an underscore is appended to the end of the filename before the extension.
                    copyfile(targetfilename, destfilename)
                    targetfilename = destfilename
                    copiedfile2 = True
                    # A copy of the Word document is made to be very sure that the original file is safe.
                doc = docx.Document(targetfilename)
                # Word document is opened as doc object.
            except:
                self.errorFileload()
                # If the Word document cannot be opened for any reason, an error box will pop up.
                return
            fulltext = [] # The fulltext list will be populated with paragraphs in the Word document.
            paragraphcount = len(doc.paragraphs)
            i=-1
            while i < paragraphcount:
                runscount = len(doc.paragraphs[i].runs)
                j=-1
                while j < runscount and runscount > 0:
                    if(doc.paragraphs[i].runs[j].font.strike == True):
                        doc.paragraphs[i].runs[j].clear()
                        # Strike-through font is eliminated from the text.
                    j+=1
                fulltext.append(doc.paragraphs[i].text) # A paragraph is added to the fulltext list.
                i+=1
            readtext = '\n'.join(fulltext) # The fulltext list is converted into one single string.
            global utf 
            utf = True # The utf value will indicate that the readtext is in utf-8 format.
        elif targetfilename[-4:len(targetfilename)]!='.txt': # If the file does not have a doc, docx, or txt extension, it will be rejected.
            self.errorInvalidfilename()
            # An error box pops up if an invalid file extension is detected.
            return
            
        if self.Function1.isChecked():
            #print('loading Function1 file')
            if copiedfile1: # If a copy of a Word 97-2003 Document was made and then converted into text file.
                renamedtargetfilename = targetfilename.replace('.txt', '')
                destinationfile = renamedtargetfilename + '_Function1.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            elif copiedfile2: # If a copy of the Word document was made using python-docx.
                renamedtargetfilename = targetfilename.replace('.docx', '')
                destinationfile = renamedtargetfilename + '_Function1.txt'
                # No need to use opentextfile() function, since readtext was already generated.
            else:
                renamedtargetfilename = targetfilename.replace('.txt', '') # If the inputted file was a text file.
                destinationfile = renamedtargetfilename + '_Function1.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            if os.access(destinationfile, os.R_OK): # Check to see if the results file already exists.
                destinationfile = destinationfile[0:-4] + '(' + filenumber(destinationfile[0:-4]) + ').txt'
                #If the results file already exists, the file name is renamed to avoid overwriting existing files.
                #For example, if SAMPLE_Function1.txt already exists, then the file is renamed SAMPLE_Function1(1).txt
                
            results = function1(readtext, destinationfile)
        elif self.Function2.isChecked():
            #print('loading Function2 file')
            if copiedfile1: # If a copy of a Word 97-2003 Document was made and then converted into text file.
                renamedtargetfilename = targetfilename.replace('.txt', '')
                destinationfile = renamedtargetfilename + '_Function2.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            elif copiedfile2: # If a copy of the Word document was made using python-docx.
                renamedtargetfilename = targetfilename.replace('.docx', '')
                destinationfile = renamedtargetfilename + '_Function2.txt'
                # No need to use opentextfile() function, since readtext was already generated.
            else:
                renamedtargetfilename = targetfilename.replace('.txt', '') # If the inputted file was a text file.
                destinationfile = renamedtargetfilename + '_Function2.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            if os.access(destinationfile, os.R_OK): # Check to see if the results file already exists.
                destinationfile = destinationfile[0:-4] + '(' + filenumber(destinationfile[0:-4]) + ').txt'
                #If the results file already exists, the file name is renamed to avoid overwriting existing files.
                #For example, if SAMPLE_Function2.txt already exists, then the file is renamed SAMPLE_Function2(1).txt
                
            results = function2(readtext, destinationfile)
        elif self.Function3.isChecked():
            #print('loading Function3 file')
            if copiedfile1: # If a copy of a Word 97-2003 Document was made and then converted into text file.
                renamedtargetfilename = targetfilename.replace('.txt', '')
                destinationfile = renamedtargetfilename + '_Function3.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            elif copiedfile2: # If a copy of the Word document was made using python-docx.
                renamedtargetfilename = targetfilename.replace('.docx', '')
                destinationfile = renamedtargetfilename + '_Function3.txt'
                # No need to use opentextfile() function, since readtext was already generated.
            else:
                renamedtargetfilename = targetfilename.replace('.txt', '') # If the inputted file was a text file.
                destinationfile = renamedtargetfilename + '_Function3.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            if os.access(destinationfile, os.R_OK): # Check to see if the results file already exists.
                destinationfile = destinationfile[0:-4] + '(' + filenumber(destinationfile[0:-4]) + ').txt'
                #If the results file already exists, the file name is renamed to avoid overwriting existing files.
                #For example, if SAMPLE_Function3.txt already exists, then the file is renamed SAMPLE_Function3(1).txt
            results = function3(readtext, destinationfile)
        elif self.Function4.isChecked():
            #print('loading Function4 file')
            if copiedfile1: # If a copy of a Word 97-2003 Document was made and then converted into text file.
                renamedtargetfilename = targetfilename.replace('.txt', '')
                destinationfile = renamedtargetfilename + '_Function4.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            elif copiedfile2: # If a copy of the Word document was made using python-docx.
                renamedtargetfilename = targetfilename.replace('.docx', '')
                destinationfile = renamedtargetfilename + '_Function4.txt'
                # No need to use opentextfile() function, since readtext was already generated.
            else:
                renamedtargetfilename = targetfilename.replace('.txt', '') # If the inputted file was a text file.
                destinationfile = renamedtargetfilename + '_Function4.txt'
                readtext = opentextfile(targetfilename) # The results text file is opened.
            if os.access(destinationfile, os.R_OK): #Checking if the results file already exists.
                destinationfile = destinationfile[0:-4] + '(' + filenumber(destinationfile[0:-4]) + ').txt'
                #If the results file already exists, the file name is renamed to avoid overwriting existing files.
                #For example, if SAMPLE_Function4.txt already exists, then the file is renamed SAMPLE_Function4(1).txt
            results = function4(readtext, destinationfile) 
        if copiedfile1: # If a copy of a Word 97-2003 Document was made and then converted into text file.
            os.remove(destfilename) # Deletes the Word copy of the file
            os.remove(targetfilename) # Deletes the Text copy of the file
        elif copiedfile2: # If a copy of the Word document was made using python-docx.
            os.remove(targetfilename) # Deletes the Word copy of the file
    
    # The following three methods set up dragging and dropping for the app
    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls:
            e.accept()
        else:
            e.ignore()

    def dragMoveEvent(self, e):
        if e.mimeData().hasUrls:
            e.accept()
        else:
            e.ignore()

    def dropEvent(self, event):
        if event.mimeData().hasUrls:
            event.setDropAction(QtCore.Qt.CopyAction)
            event.accept()
            self.fname = []
            for url in event.mimeData().urls():
                self.fname.append(str(url.toLocalFile()))
            self.load_file()
        else:
            event.ignore()
            
    # The following are the functions for the warning boxes that come up upon triggering an error.
 
         
    def errorFileload(self):
         messagetext = 'File loading error.'
         QtWidgets.QMessageBox.information(self, "Warning", messagetext)
         return
         
    def errorInvalidfilename(self):
         messagetext = 'Invalid file.'
         QtWidgets.QMessageBox.information(self, "Warning", messagetext)
         return

    def errorpywin32(self):
         messagetext = 'Microsoft Word COM object not identified in registry.'
         QtWidgets.QMessageBox.information(self, "Information", messagetext)
         return

app = QtWidgets.QApplication(sys.argv)
form = MainApplication()
form.show()
app.exec_()

